# This is a sample Python script.

# Press Shift+F10 to execute it or replace it with your code.
# Press Double Shift to search everywhere for classes, files, tool windows, actions, and settings.

import os
import sys, requests

import csv
import json

import openpyxl as xl
from openpyxl.chart import BarChart, Reference

import docx
import PyPDF2
from collections import OrderedDict
from PyPDF2 import PdfFileReader

def print_hi(name):
    # Use a breakpoint in the code line below to debug your script.
    print(f'Hi, {name}')  # Press Ctrl+F8 to toggle the breakpoint.

def change_file_name():
    os.chdir('C:\\Users\\Esther\\Desktop\\example')
    for f in os.listdir():
        f_name, f_ext = os.path.splitext(f)
        f_title, f_course, f_num = f_name.split('-')
        #elining all names
        f_title = f_title.strip()
        f_course = f_course.strip()
        f_num = f_num.strip()[1:].zfill(2)# ignor the num # + filling zeros before number
        print('{}-{}-{}{}'.format(f_num, f_course, f_title, f_ext))
        new_name = '{}-{}{}'.format(f_num, f_title, f_ext)

        os.rename(f, new_name)

def excel_handle():

    wb = xl.load_workbook('exm1.xlsx')
    sheet = wb['Sheet1']
    cell1 = sheet['a1']
    cell2 = sheet.cell(1,1) #same as cell1

    for row in range(2, sheet.max_row + 1):
        cell = sheet.cell(row, 3)
        val = cell.value[1:]
        corrected_price = float(val) * 0.9
        corrected_price_cell = sheet.cell(row, 4)
        corrected_price_cell.value = corrected_price

    last_cell = sheet.cell(1,sheet.max_row).value = 'new_price'
    #selecting the needed args
    values = Reference(sheet,
              min_row=2,
              max_row=sheet.max_row,
              min_col=4,
              max_col=4)

    chart = BarChart()
    chart.add_data(values)
    sheet.add_chart(chart, 'e2')

    wb.save('exm2.xlsx')


def word_handle(file):
    doc = docx.Document(file)
    completedText = []

    for p in doc.paragraphs:
        completedText.append(p.text)

    return '\n' .join(completedText)


#read the fields within any PDF form doc
def _getFields(obj, tree=None, retval=None, fileobj=None):
    fieldAttributes = {'/FT': 'Field Type', '/Parent': 'Parent', '/T': 'Field Name',
    '/TU': 'Alternate Field Name', '/TM': 'Mapping Name', '/Ff': 'Field Flags',
    '/V': 'Value', '/DV': 'Default Value'}
    if retval is None:
        retval = OrderedDict()
        catalog = obj.trailer["/Root"]
        if "/AcroForm" in catalog:
            tree = catalog["/AcroForm"]
        else:
            return None
    if tree is None:
        return retval

    obj._checkKids(tree, retval, fileobj)
    for attr in fieldAttributes:
        if attr in tree:
            obj._buildField(tree, retval, fileobj, fieldAttributes)
            break

    if "/Fields" in tree:
        fields = tree["/Fields"]
        for f in fields:
            field = f.getObject()
            obj._buildField(field, retval, fileobj, fieldAttributes)

    return retval

#reads the PDF form and returns all the fields vals contained in PDF file
def get_form_fields(infile):
    infile = PdfFileReader(open(infile, 'rb'))
    fields = _getFields(infile)
    return OrderedDict((k, v.get('/V', '')) for k, v in fields.items())

#JavaScript func that is capable od selecting at runtime
def selectListOption(all_lines, k, v):
    all_lines.append('function setSelectedIndex(s, v) {')
    all_lines.append('for (var i = 0; i < s.options.length; i++) {')
    all_lines.append('if (s.options[i].text == v) {')
    all_lines.append('s.options[i].selected = true;')
    all_lines.append('return;')
    all_lines.append('}')
    all_lines.append('}')
    all_lines.append('}')
    all_lines.append('setSelectedIndex(document.getElementById("' + k + '"), "' + v + '");')


def readList(fname):
    lst = []
    with open(fname, 'r') as fh:
        for l in fh:
            lst.append(l.rstrip(os.linesep))
    return lst

#responsible for creating the JavaScript script that will be executed on the browser.
def createBrowserScript(fl, fl_ext, items, pdf_file_name):
    if pdf_file_name and len(fl) > 0:
        of = os.path.splitext(pdf_file_name)[0] + '.txt'
        all_lines = []
        for k, v in items.items():
            print(k + ' -> ' + v)
            if (v in ['/Yes', '/On']):
                all_lines.append("document.getElementById('" + k + "').checked = true;\n");
            elif (v in ['/0'] and k in fl_ext):
                all_lines.append("document.getElementById('" + k + "').checked = true;\n");
            elif (v in ['/No', '/Off', '']):
                all_lines.append("document.getElementById('" + k + "').checked = false;\n");
            elif (v in [''] and k in fl_ext):
                all_lines.append("document.getElementById('" + k + "').checked = false;\n");
            elif (k in fl):
                selectListOption(all_lines, k, v)
            else:
                all_lines.append("document.getElementById('" + k + "').value = '" + v + "';\n");
        outF = open(of, 'w')
        outF.writelines(all_lines)
        outF.close()

def execute(args):
    try:
        fl = readList('myview.ini')
        fl_ext = readList('myview_ext.ini')
        if len(args) == 2:
            pdf_file_name = args[1]
            items = get_form_fields(pdf_file_name)
            createBrowserScript(fl, fl_ext, items, pdf_file_name)
        else:
            files = [f for f in os.listdir('.') if os.path.isfile(f) and f.endswith('.pdf')]
            for f in files:
                items = get_form_fields(f)
                createBrowserScript(fl, fl_ext, items, f)
    except BaseException as msg:
        print('An error occurred... :( ' + str(msg))

def pdf_handle():
    pdf1File = open('meetingminutes.pdf', 'rb') #rb = read binary object contains our file
    pdf2File = open('meetingminutes2.pdf', 'rb')
    pdf1Reader = PyPDF2.PdfFileReader(pdf1File)
    pdf2Reader = PyPDF2.PdfFileReader(pdf2File)
    pdfWriter = PyPDF2.PdfFileWriter()

    for pNum in range(pdf1Reader.numPages):
        pObj = pdf1Reader.getPage(pNum)
        pdfWriter.addPage(pObj)

    for pNum in range(pdf2Reader.numPages):
        pObj = pdf2Reader.getPage(pNum)
        pdfWriter.addPage(pObj)

    pdfOutputFile = open('combinedminutes.pdf', 'wb')
    pdfWriter.write(pdfOutputFile)
    pdfOutputFile.close()
    pdf1File.close()
    pdf2File.close()

def word_handle2(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        #fullText.append(para.text)
        fullText.append(' ' + para.text)
    return '\n'.join(fullText)

def cvs_handle(filename):
    os.makedirs('headerRemoved', exist_ok=True)
    # Loop through every file in the current working directory.
    for csvFilename in os.listdir('.'):
        if not csvFilename.endswith('.csv'):
            continue  # skip non-csv files
        print('Removing header from ' + csvFilename + '...')

        # Read the CSV file in (skipping first row).
        csvRows = []
        csvFileObj = open(csvFilename)
        readerObj = csv.reader(csvFileObj)
        for row in readerObj:
            if readerObj.line_num == 1:
                continue  # skip first row
            csvRows.append(row)
        csvFileObj.close()

        # Write out the CSV file.
        csvFileObj = open(os.path.join('headerRemoved', csvFilename), 'w',
                          newline='')
        csvWriter = csv.writer(csvFileObj)
        for row in csvRows:
            csvWriter.writerow(row)
        csvFileObj.close()

def json_handle():
    # Compute location from command line arguments.
    if len(sys.argv) < 2:
        print('Usage: quickWeather.py location')
        sys.exit()
    location = ' '.join(sys.argv[1:])

    # Download the JSON data from OpenWeatherMap.org's API.
    url = 'https://openweathermap.org/' % (location)
    response = requests.get(url)
    response.raise_for_status()
    # Load JSON data into a Python variable.
    weatherData = json.loads(response.text)
    # Print weather descriptions.
    w = weatherData['list']
    print('Current weather in %s:' % (location))
    print(w[0]['weather'][0]['main'], '-', w[0]['weather'][0]['description'])
    print()
    print('Tomorrow:')
    print(w[1]['weather'][0]['main'], '-', w[1]['weather'][0]['description'])
    print()
    print('Day after tomorrow:')
    print(w[2]['weather'][0]['main'], '-', w[2]['weather'][0]['description'])


# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    print_hi('PyCharm')
    os.chdir('C:\\Users\\Esther\\Desktop\\example')
    #change_file_name()
    #excel_handle()
    print(word_handle2('file1.docx'))
    #pdf_handle()
    json_handle()

# See PyCharm help at https://www.jetbrains.com/help/pycharm/
